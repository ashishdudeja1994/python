# Import docx NOT python-docx
import os,sys,glob
import docx
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from docx2pdf import convert
# from pdf2image import convert_from_path
from pdf2jpg import pdf2jpg
import numpy as np
commands = ['st cell', 'alt', 'hget retsubunit operationalState|userlabel|electricalant|maxtilt|mintilt|iuantbasestationid|iuantsectorid', 'mfitr']
doc = docx.Document()
for file1 in glob.glob("./input_files/*.txt"):
    with open(file1,'r') as f:
        # all_data = f.read()
        all_lines_list = f.readlines()
        sitename = [x[:x.find(">")] for x in all_lines_list if x.find('st cell') != -1][0]
        for command_index in range(0,len(commands)):
            start_line = None
            for line_index in range(0,len(all_lines_list)):
                if all_lines_list[line_index].find(commands[command_index]) != -1:
                    start_line = line_index
                elif all_lines_list[line_index].find(sitename) != -1 and start_line is not None:
                    end_line = line_index
                    break
            # print(item)
            # doc.add_heading('GeeksForGeeks', 0)
            doc.add_paragraph().add_run("".join(all_lines_list[start_line:end_line-1]))
            if command_index != len(commands)-1:
                doc.add_page_break()
        shd = OxmlElement('w:background')

        # # Add attributes to the element
        shd.set(qn('w:color'), '0D0D0D')
        shd.set(qn('w:themeColor'), 'text1')
        shd.set(qn('w:themeTint'), 'F2')
        doc.element.insert(0,shd)
        # print(doc.element.xml[1000:1500])
        shd1 = OxmlElement('w:displayBackgroundShape')
        doc.settings.element.insert(0,shd1)

        # # Add attributes to the element
        # print(doc.settings.element.xml)
        # print(highlight_para.style._element.xpath("./w:pPr/w:pBdr/w:bottom"))
        # from docx.oxml.ns import qn

        # bottom = doc.styles.element.xpath("./w:background")[0]
        # bottom.set(qn("w:color"), "FF00FF")
        # # Now save the document to a location 
        # doc.save('./output_files/' + os.path.basename(file1).split(".")[0] + '.docx')
        # convert('./output_files/' + os.path.basename(file1).split(".")[0] + '.docx', './output_files/' + os.path.basename(file1).split(".")[0] + '.pdf')
        # result = pdf2jpg.convert_pdf2jpg('./output_files/' + os.path.basename(file1).split(".")[0] + '.pdf', "./output_files/images/", pages="ALL")
        import cv2

        img = cv2.imread(r"C:\Users\ashis\Documents\Work\python\python_mohit\output_files\images\log (17).pdf_dir\0_log (17).pdf.jpg") # Read in the image and convert to grayscale
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2WHITE)
        gray = 255*(gray < 128).astype(np.uint8) # To invert the text to white
        coords = cv2.findNonZero(gray) # Find all non-zero points (text)
        x, y, w, h = cv2.boundingRect(coords) # Find minimum spanning bounding box
        rect = img[y:y+h, x:x+w] # Crop the image - note we do this on the original image
        # cv2.imshow("Cropped", rect) # Show it
        # cv2.waitKey(0)
        # cv2.destroyAllWindows()
        cv2.imwrite("rect.png", rect) # Save the image
        # pages = convert_from_path('./output_files/' + os.path.basename(file1).split(".")[0] + '.pdf', 500)
        # image_count = 1
        # for page in pages:
        #     page.save('./output_files/' + os.path.basename(file1).split(".")[0] + '_' + str(image_count) + '.jpg', 'JPEG')
        #     image_count += 1
# sys.exit()
# Create an instance of a word document

  
# Add a Title to the document 
# 
  
# Creating paragraph with some content and Highlighting it.
# with open('./text1.txt') as f:
#     highlight_para = doc.add_paragraph().add_run(f.read())
#     paragraphs = doc.paragraphs
# print(paragraphs)
# print(type(doc.element.xml))
# xml_string = doc.element.xml
# print(doc.element.append("w:background"))
# for item in doc.element.getchildren():
#     print(item)
# Create XML element


# document = Document()
# >>> run = document.add_heading(u'', 0).add_run('hello world')
# >>> paragraphs = document.paragraphs
# >>> print(paragraphs[0].style._element.xml)